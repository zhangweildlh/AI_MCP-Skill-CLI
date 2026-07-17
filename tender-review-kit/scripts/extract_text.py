#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""extract_text.py — 统一取数：PDF / Word(.docx) → 「带行号文本 + 结构化表格」

tender-review-skill 的地基程序：把任意格式的招标文件转成统一的、带行号的纯文本，
供后续判决词扫描 / 查漏 / 各 subagent 通读。
行号是提取后「自编的定位锚点」（文件本身没有行号，PDF/Word 都没有）。

用法:
    python extract_text.py <file.docx|file.pdf> [--outdir DIR]

输出:
    <outdir>/<stem>.lines.txt    每行: 行号<TAB>文本（正文段落 + 表格摊平，统一连续编号）
    <outdir>/<stem>.tables.json  结构化表格 [{table_id,line_start,line_end,n_rows,n_cols,rows}]（仅 docx）

支持:   .docx (python-docx) / .pdf (pdftotext -layout，回退 pypdf)
不支持: .doc (老 Word) → 提示另存为 .docx 或 .pdf
"""
import sys
import json
import subprocess
import argparse
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


def extract_docx(path):
    """按文档顺序遍历段落和表格。

    python-docx 默认把 doc.paragraphs 和 doc.tables 分开返回，丢失原文顺序；
    这里手动遍历 body 的子元素，保持正文与表格的真实先后。
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = Document(str(path))
    lines, tables = [], []
    tbl_id = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            txt = Paragraph(child, doc).text.strip()
            if txt:
                lines.append(txt)
        elif child.tag == qn("w:tbl"):
            tbl_id += 1
            t = Table(child, doc)
            start = len(lines) + 1
            rows = []
            for row in t.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(cells)
                lines.append("[T%d] %s" % (tbl_id, " | ".join(cells)))
            tables.append({
                "table_id": tbl_id,
                "line_start": start,
                "line_end": len(lines),
                "n_rows": len(rows),
                "n_cols": max((len(r) for r in rows), default=0),
                "rows": rows,
            })
    return lines, tables


def extract_pdf(path):
    """优先 pdftotext -layout（保留版面，表格对齐好）；失败或空输出则回退 pypdf。"""
    text = ""
    try:
        out = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            capture_output=True,
            check=True,
        )
        text = out.stdout.decode("utf-8", errors="replace")
        if not text.strip():
            raise RuntimeError("pdftotext empty output")
    except Exception:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
    lines = [ln.rstrip() for ln in text.splitlines()]
    return lines, []


def main():
    ap = argparse.ArgumentParser(description="统一取数：PDF/Word → 带行号文本 + 结构化表格")
    ap.add_argument("file")
    ap.add_argument("--outdir", default=None, help="输出目录（默认 <文件目录>/_extracted）")
    ap.add_argument("--name", default=None,
                    help="输出文件名简称(覆盖默认的源文件名,避免长文件名衍生一堆)")
    args = ap.parse_args()

    path = Path(args.file)
    if not path.exists():
        print("ERROR file_not_found:", path)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext == ".docx":
        lines, tables = extract_docx(path)
    elif ext == ".pdf":
        lines, tables = extract_pdf(path)
    elif ext == ".doc":
        print("ERROR doc_unsupported: .doc(老 Word) 暂不支持，请在 Word 里另存为 .docx 或 .pdf 后重试")
        sys.exit(2)
    else:
        print("ERROR bad_format:", ext, "(仅支持 .docx / .pdf)")
        sys.exit(2)

    outdir = Path(args.outdir) if args.outdir else path.parent / "_extracted"
    outdir.mkdir(parents=True, exist_ok=True)
    stem = args.name if args.name else path.stem

    nonempty = sum(1 for ln in lines if ln.strip())
    chars = sum(len(ln) for ln in lines)

    # 质量门槛(先检查后写)：扫描件/图片型 PDF、加密或提取异常会得到空/极少文本。
    # 失败时**不落盘**——否则会留下垃圾 lines.txt + 后续撒网扫空、静默出"没有废标点"的假报告。
    MIN_CHARS = 50
    if nonempty == 0 or chars < MIN_CHARS:
        print("ERROR empty_or_low_extraction format=%s nonempty=%d chars=%d"
              % (ext, nonempty, chars))
        print("  提取内容为空或过少——常见原因：扫描件/图片型 PDF(需 OCR)、加密 PDF、或提取异常。")
        print("  已中止,未写出任何文件;请人工核对原文件,勿进入后续撒网扫描。")
        sys.exit(3)

    lines_path = outdir / (stem + ".lines.txt")
    with open(lines_path, "w", encoding="utf-8") as f:
        for i, ln in enumerate(lines, 1):
            f.write("%d\t%s\n" % (i, ln))

    tables_path = None
    if tables:
        tables_path = outdir / (stem + ".tables.json")
        with open(tables_path, "w", encoding="utf-8") as f:
            json.dump(tables, f, ensure_ascii=False, indent=1)

    print("OK format=%s" % ext)
    print("lines_total=%d lines_nonempty=%d chars=%d tables=%d"
          % (len(lines), nonempty, chars, len(tables)))
    print("out_lines=%s" % lines_path)
    if tables_path:
        print("out_tables=%s" % tables_path)


if __name__ == "__main__":
    main()
