#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""generate_fixture.py — 生成回归测试用的合成招标文件(.docx)

合成一份纯虚构的"DEMO 招标文件",含已知数量的:
  - 一级判决词命中(否决/无效/视为/废标/拒收 等)
  - 关系门槛(加盖原厂公章/必须接入现有平台/与已有系统兼容 等)
  - ▲ 标识参数(技术规格)
  - 证明文件要求(检测报告/CCC 认证 等)

跑一次:`python tests/generate_fixture.py` → tests/fixtures/sample_tender.docx
test_smoke.py 用它做回归基准,不依赖任何真实标书。

合成内容**完全虚构**(公司名/数字/参数全是占位符),开源仓库可以放心带这份文件。
"""
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document


def build():
    doc = Document()
    # 封面 + 项目元信息
    doc.add_heading("DEMO 设备采购招标文件", 0)
    doc.add_paragraph("招标编号:DEMO-2024-001")
    doc.add_paragraph("招标人:虚构科技有限公司(示例)")
    doc.add_paragraph("招标机构:示例招标代理机构")

    # 第一章 招标公告
    doc.add_heading("第一章 招标公告", 1)
    doc.add_paragraph("1.1 招标内容:本项目采购 X 设备 50 套,Y 模块 100 个。")
    doc.add_paragraph("1.2 投标人资格:")
    doc.add_paragraph("  1.2.1 申请人应为境内注册独立法人,具有独立承担民事责任的能力。")
    doc.add_paragraph("  1.2.2 不接受联合体投标。")
    doc.add_paragraph("  1.2.3 投标人有重大失信记录的,其投标予以否决。")

    # 第二章 投标人须知
    doc.add_heading("第二章 投标人须知", 1)
    doc.add_paragraph("2.1 投标报价不得超过最高投标限价,否则按无效投标处理。")
    doc.add_paragraph("2.2 投标保证金 5 万元,未按时按额提交的,其投标将被否决。")
    doc.add_paragraph("2.3 投标文件未签字或未盖章的,招标人予以拒收。")
    doc.add_paragraph("2.4 投标文件密封不符合要求或逾期送达的,招标人予以拒收。")
    doc.add_paragraph("2.5 实质性偏离招标文件要求的,投标将被否决。")
    doc.add_paragraph("2.6 业绩材料缺失或业主签盖不齐的,视为无效业绩。")
    doc.add_paragraph("2.7 中标人不能按要求递交履约保证金的,视为放弃中标。")
    doc.add_paragraph("2.8 投标人有任何弄虚作假行为,其投标将被否决。")

    # 第三章 评标办法
    doc.add_heading("第三章 评标办法(综合评分法)", 1)
    doc.add_paragraph("3.1 总分 100 分 = 价格分 40 + 技术分 40 + 商务分 20。")
    doc.add_paragraph("3.2 价格分:最低价得 40 分,其他按 基准价/报价 ×40 计算。")
    doc.add_paragraph("3.3 商务分明细")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "评分项"
    tbl.rows[0].cells[1].text = "分值"
    tbl.rows[0].cells[2].text = "评分梯度"
    tbl.rows[1].cells[0].text = "业绩"
    tbl.rows[1].cells[1].text = "8 分"
    tbl.rows[1].cells[2].text = "近三年类似业绩 ≥3 项得满分,每减少 1 项扣 2 分"
    tbl.rows[2].cells[0].text = "资质"
    tbl.rows[2].cells[1].text = "5 分"
    tbl.rows[2].cells[2].text = "具有 ISO9001 得 3 分,具有 ISO14001 加 2 分,无认证不得分"
    tbl.rows[3].cells[0].text = "财务"
    tbl.rows[3].cells[1].text = "4 分"
    tbl.rows[3].cells[2].text = "近三年盈利得 4 分,亏损不得分"
    tbl.rows[4].cells[0].text = "信誉"
    tbl.rows[4].cells[1].text = "3 分"
    tbl.rows[4].cells[2].text = "近三年无失信记录得 3 分,否则得 0 分"
    doc.add_paragraph("3.4 投标文件未对实质性要求作响应,投标将被否决。")

    # 第四章 技术规范(▲ 标识参数)
    doc.add_heading("第四章 技术规范", 1)
    doc.add_paragraph("4.1 X 设备技术参数(▲ 项为实质性要求,不满足即否决):")
    doc.add_paragraph("  ▲ 4.1.1 工作频率 ≥ 2.4GHz")
    doc.add_paragraph("  ▲ 4.1.2 输入电压 100~240V AC")
    doc.add_paragraph("  ▲ 4.1.3 工作温度 -20℃ ~ +60℃")
    doc.add_paragraph("  ▲ 4.1.4 防护等级 ≥ IP65,提供检测报告")
    doc.add_paragraph("  ▲ 4.1.5 平均无故障时间 MTBF > 10 万小时,提供检测报告")
    doc.add_paragraph("  ▲ 4.1.6 接口 ≥4 路 HDMI")
    doc.add_paragraph("  ▲ 4.1.7 支持 H.264 / H.265 编解码")
    doc.add_paragraph("  ▲ 4.1.8 待机功耗 ≤ 5W")
    doc.add_paragraph("  ▲ 4.1.9 具有 CCC 认证证书,加盖原厂公章")
    doc.add_paragraph("  ▲ 4.1.10 具有 ISO9001 认证证书")
    doc.add_paragraph("  ▲ 4.1.11 必须接入现有 X 平台,提供原厂授权")
    doc.add_paragraph("  ▲ 4.1.12 与已有 Y 系统兼容,需提供兼容性测试报告")

    doc.add_paragraph("4.2 关系门槛:")
    doc.add_paragraph("  4.2.1 投标人须为指定品牌的原厂或其授权代理。")
    doc.add_paragraph("  4.2.2 第三方软件须加盖原厂公章。")
    doc.add_paragraph("  4.2.3 X 设备必须与采购方现有平台无缝对接。")

    # 第五章 证明文件
    doc.add_heading("第五章 证明文件清单", 1)
    doc.add_paragraph("投标人须提供:")
    doc.add_paragraph("  5.1 营业执照、组织机构代码证、税务登记证")
    doc.add_paragraph("  5.2 提供 CCC 认证证书复印件")
    doc.add_paragraph("  5.3 提供检测报告(CMA 或 CNAS 资质)")
    doc.add_paragraph("  5.4 提供 ISO9001 认证证书")
    doc.add_paragraph("  5.5 提供入网证书复印件")
    doc.add_paragraph("  5.6 提供制造商出具的授权证明")

    out = Path(__file__).resolve().parent / "fixtures" / "sample_tender.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print("OK saved:", out)
    return out


if __name__ == "__main__":
    build()
